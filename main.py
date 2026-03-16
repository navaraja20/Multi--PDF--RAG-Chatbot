from __future__ import annotations

import logging
import sys
from pathlib import Path

from tqdm import tqdm
from docx import Document

from scraper.collect_urls import load_urls
from scraper.clean_text import clean_text
from scraper.convert_to_pdf import docx_to_pdf
from scraper.scrape_pages import (
    DEFAULT_HEADERS,
    extract_text_and_title,
    fetch_html,
    iter_urls,
    rate_limit_sleep,
    sanitize_title,
)

BASE_DIR = Path(__file__).parent
URL_FILE = BASE_DIR / "urls" / "urls.txt"
DOCX_DIR = BASE_DIR / "data" / "docx_files"
PDF_DIR = BASE_DIR / "data" / "pdf_files"
LOG_LEVEL = logging.INFO


def ensure_dirs() -> None:
    for path in (DOCX_DIR, PDF_DIR):
        path.mkdir(parents=True, exist_ok=True)


def setup_logging() -> None:
    logging.basicConfig(
        level=LOG_LEVEL,
        format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
        handlers=[logging.StreamHandler(sys.stdout)],
    )


def _unique_path(path: Path) -> Path:
    """Return a unique path by appending a counter if needed."""
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    parent = path.parent
    counter = 1
    while True:
        candidate = parent / f"{stem}_{counter}{suffix}"
        if not candidate.exists():
            return candidate
        counter += 1


def save_docx(cleaned_text: str, title: str, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_title = sanitize_title(title)
    doc_path = output_dir / f"{safe_title}.docx"
    doc = Document()
    doc.add_heading(title, level=0)
    for paragraph in cleaned_text.split("\n\n"):
        doc.add_paragraph(paragraph)
    doc.save(doc_path)
    return doc_path


def process_url(url: str, session) -> None:
    html = fetch_html(url, session=session)
    if not html:
        return
    raw_text, title = extract_text_and_title(html, url=url)
    cleaned = clean_text(raw_text)
    if not cleaned.strip():
        logging.info("Skipping empty content for %s", url)
        return

    safe_title = sanitize_title(title)
    docx_path = DOCX_DIR / f"{safe_title}.docx"
    pdf_target = PDF_DIR / f"{safe_title}.pdf"

    if docx_path.exists() and pdf_target.exists():
        logging.info("Skipping existing outputs for %s -> %s", url, safe_title)
        return

    if not docx_path.exists():
        docx_path = save_docx(cleaned, title, DOCX_DIR)

    try:
        docx_to_pdf(docx_path, pdf_target)
    except Exception as exc:
        logging.warning("PDF conversion failed for %s: %s", docx_path, exc)


def main() -> None:
    setup_logging()
    ensure_dirs()
    try:
        urls = load_urls(URL_FILE)
    except FileNotFoundError as exc:
        logging.error("Cannot start: %s", exc)
        return

    if not urls:
        logging.error("No URLs found in %s", URL_FILE)
        return

    session = __import__("requests").Session()
    session.headers.update(DEFAULT_HEADERS)

    for url in tqdm(list(iter_urls(urls)), desc="Scraping", unit="page"):
        try:
            process_url(url, session=session)
        except Exception as exc:
            logging.warning("Processing failed for %s: %s", url, exc)
        finally:
            rate_limit_sleep()


if __name__ == "__main__":
    main()
